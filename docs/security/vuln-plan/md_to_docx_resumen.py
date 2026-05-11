"""Genera RESUMEN_TRABAJO_REALIZADO.docx desde el Markdown (sin pandoc)."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document


def strip_md_cells(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(line: str) -> bool:
    parts = parse_table_row(line)
    return bool(parts) and all(re.fullmatch(r":?-{3,}:?", p) is not None for p in parts)


def add_runs_with_emphasis(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def try_consume_table(lines: list[str], i: int) -> tuple[list[list[str]], int] | None:
    if i >= len(lines) or not lines[i].strip().startswith("|"):
        return None
    header = parse_table_row(lines[i])
    i += 1
    if i >= len(lines) or not is_separator_row(lines[i]):
        return None
    i += 1
    rows = [header]
    while i < len(lines):
        ln = lines[i].strip()
        if not ln.startswith("|"):
            break
        rows.append(parse_table_row(lines[i]))
        i += 1
    return rows, i


def main() -> None:
    md_path = Path(__file__).resolve().parent / "RESUMEN_TRABAJO_REALIZADO.md"
    out_path = md_path.with_suffix(".docx")
    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        if stripped == "---":
            i += 1
            continue
        if stripped == "":
            i += 1
            continue

        table_result = try_consume_table(lines, i)
        if table_result:
            tbl_rows, i = table_result
            ncol = max(len(r) for r in tbl_rows)
            table = doc.add_table(rows=len(tbl_rows), cols=ncol)
            table.style = "Table Grid"
            for ri, row in enumerate(tbl_rows):
                for ci in range(ncol):
                    cell_text = strip_md_cells(row[ci]) if ci < len(row) else ""
                    table.rows[ri].cells[ci].text = cell_text
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif re.match(r"^\d+\.\s+", stripped):
            body = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number")
            add_runs_with_emphasis(p, body)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_emphasis(p, stripped[2:])
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(stripped.strip("* ").strip())
            r.italic = True
        else:
            p = doc.add_paragraph()
            add_runs_with_emphasis(p, lines[i].rstrip())

        i += 1

    doc.save(out_path)
    print(out_path)


if __name__ == "__main__":
    main()
